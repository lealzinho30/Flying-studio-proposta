"""Gera o DOCX da proposta no formato profissional Flying Studio.

Por padrão usa o papel timbrado `papel_timbrado/TIMBRADO_FLYINGSTUDIO.docx`
(cabeçalho/rodapé do modelo; corpo gerado pela aplicação). Sem timbrado: header/footer programáticos.

Layout CLI (capa) / web (Anexo I em docx_gen.js):
- Capa com título grande, caixa roxa de cliente/projeto, investimento destacado
- Páginas seguintes: tabelas estilizadas (header roxo, zebra), considerações, assinatura

Paleta:
  primária roxa:   #7C5CFF
  primária dark:   #5B3CFF
  acento verde:    #9DDB1A
  texto:           #1F2330
  texto soft:      #5C6473
  cinza light:     #E7E9EE
  off-white:       #F7F8FB
"""
from __future__ import annotations

import datetime as _dt
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt, RGBColor, Emu

from .orcamento import Orcamento

LOGO_PATH = Path(__file__).resolve().parent / "flying_logo.png"
_PAPEL_DIR = Path(__file__).resolve().parent / "papel_timbrado"
TIMBRADO_DOCX = _PAPEL_DIR / "TIMBRADO_FLYINGSTUDIO.docx"
TIMBRADO_DOC = _PAPEL_DIR / "TIMBRADO_FLYINGSTUDIO.doc"

MESES_PT = [
    "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
    "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
]

# Paleta como tupla (R,G,B) — usada com RGBColor
COR_PRIMARIA = RGBColor(0x7C, 0x5C, 0xFF)
COR_PRIMARIA_DARK = RGBColor(0x5B, 0x3C, 0xFF)
COR_ACENTO = RGBColor(0x9D, 0xDB, 0x1A)
COR_TEXTO = RGBColor(0x1F, 0x23, 0x30)
COR_TEXTO_SOFT = RGBColor(0x5C, 0x64, 0x73)
COR_BRANCO = RGBColor(0xFF, 0xFF, 0xFF)

# Hex para shading (sem #)
HEX_PRIMARIA = "7C5CFF"
HEX_PRIMARIA_DARK = "5B3CFF"
HEX_OFFWHITE = "F7F8FB"
HEX_BRANCO = "FFFFFF"

FONTE = "Calibri"


def _brl(valor: float) -> str:
    s = f"{valor:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R${s}"


def _data_extenso(data: _dt.date) -> str:
    return f"{data.day:02d} de {MESES_PT[data.month - 1]} de {data.year}"


_UNIDADES = ["", "Um", "Dois", "Três", "Quatro", "Cinco", "Seis", "Sete", "Oito", "Nove",
             "Dez", "Onze", "Doze", "Treze", "Quatorze", "Quinze", "Dezesseis", "Dezessete",
             "Dezoito", "Dezenove"]
_DEZENAS = ["", "", "Vinte", "Trinta", "Quarenta", "Cinquenta", "Sessenta", "Setenta", "Oitenta", "Noventa"]
_CENTENAS = ["", "Cento", "Duzentos", "Trezentos", "Quatrocentos", "Quinhentos",
             "Seiscentos", "Setecentos", "Oitocentos", "Novecentos"]


def _ate_999(n: int) -> str:
    if n == 0:
        return ""
    if n == 100:
        return "Cem"
    partes: list[str] = []
    c, r = divmod(n, 100)
    if c:
        partes.append(_CENTENAS[c])
    if r < 20:
        if r:
            partes.append(_UNIDADES[r])
    else:
        d, u = divmod(r, 10)
        if u:
            partes.append(f"{_DEZENAS[d]} e {_UNIDADES[u]}")
        else:
            partes.append(_DEZENAS[d])
    return " e ".join(partes)


def _extenso(valor: float) -> str:
    inteiro = int(round(valor))
    if inteiro == 0:
        return "Zero Reais"

    milhoes, resto = divmod(inteiro, 1_000_000)
    milhares, unidades = divmod(resto, 1_000)

    blocos: list[str] = []
    if milhoes:
        blocos.append(f"{_ate_999(milhoes)} {'Milhão' if milhoes == 1 else 'Milhões'}")
    if milhares:
        if milhares == 1:
            blocos.append("Mil")
        else:
            blocos.append(f"{_ate_999(milhares)} Mil")
    if unidades:
        blocos.append(_ate_999(unidades))

    texto = ", ".join(blocos) if len(blocos) > 1 else blocos[0]
    return f"{texto} {'Real' if inteiro == 1 else 'Reais'}"


# ---------- helpers de baixo nível para shading e bordas ----------


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, *, bottom_color="E7E9EE", bottom_size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    # remove bordas existentes
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(bottom_size))
    bot.set(qn("w:color"), bottom_color)
    borders.append(bot)
    tc_pr.append(borders)


def _set_cell_no_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right", "bottom"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tc_pr.append(borders)


def _add_horizontal_line(paragraph, *, color="7C5CFF", size=8):
    """Adiciona uma borda inferior fina no parágrafo (linha horizontal)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    p_bdr.append(bot)
    p_pr.append(p_bdr)


# ---------- estilo de runs/parágrafos ----------


def _run(p, texto, *, bold=False, italic=False, size=11, color=COR_TEXTO):
    r = p.add_run(texto)
    r.font.name = FONTE
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _add_par(doc, texto="", *, bold=False, italic=False, size=11, color=COR_TEXTO,
             alignment=None, space_after=4, space_before=0):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.25
    if texto:
        _run(p, texto, bold=bold, italic=italic, size=size, color=color)
    return p


def _add_titulo_secao(doc, numero, titulo):
    """01.\nAPRESENTAÇÃO   (numerinho roxo + título grande preto)"""
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(8)
    _run(p1, f"{numero}.", bold=True, size=9, color=COR_PRIMARIA)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    _run(p2, titulo.upper(), bold=True, size=18, color=COR_TEXTO)
    return p2


def _add_bullet(doc, texto, *, label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    _run(p, "•  ", bold=True, color=COR_PRIMARIA, size=11)
    if label:
        _run(p, f"{label}: ", bold=True, color=COR_TEXTO, size=11)
    _run(p, texto, color=COR_TEXTO, size=11)
    return p


# ---------- papel timbrado ----------


def _converter_doc_para_docx(doc_path: Path) -> Path:
    lo = shutil.which("libreoffice") or shutil.which("soffice")
    if not lo:
        raise FileNotFoundError(
            f"Arquivo .doc encontrado ({doc_path}), mas LibreOffice não está instalado. "
            "Salve como .docx no Word ou defina FLYING_TIMBRADO apontando para o .docx."
        )
    tmp = Path(tempfile.mkdtemp())
    subprocess.run(
        [lo, "--headless", "--convert-to", "docx", "--outdir", str(tmp), str(doc_path)],
        check=True,
        capture_output=True,
    )
    out = tmp / f"{doc_path.stem}.docx"
    if not out.exists():
        raise FileNotFoundError(f"Conversão falhou: {doc_path}")
    dest = TIMBRADO_DOCX
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(out, dest)
    shutil.rmtree(tmp, ignore_errors=True)
    return dest


def resolver_timbrado(explicito: Path | None = None) -> Path | None:
    candidatos: list[Path] = []
    if explicito is not None:
        candidatos.append(explicito)
    env = os.getenv("FLYING_TIMBRADO")
    if env:
        candidatos.append(Path(env))
    candidatos.extend([
        TIMBRADO_DOCX,
        TIMBRADO_DOC,
        Path(r"O:/PAPEL_TIMBRADO/TIMBRADO_FLYINGSTUDIO.docx"),
        Path(r"O:/PAPEL_TIMBRADO/TIMBRADO_FLYINGSTUDIO.doc"),
        Path("/PAPEL_TIMBRADO/TIMBRADO_FLYINGSTUDIO.docx"),
        Path("/mnt/o/PAPEL_TIMBRADO/TIMBRADO_FLYINGSTUDIO.docx"),
        Path("/mnt/o/PAPEL_TIMBRADO/TIMBRADO_FLYINGSTUDIO.doc"),
    ])
    for p in candidatos:
        if not p.exists():
            continue
        if p.suffix.lower() == ".doc":
            return _converter_doc_para_docx(p)
        return p
    return None


def _limpar_corpo(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


# ---------- header / footer ----------


def _setup_header(doc):
    section = doc.sections[0]
    header = section.header
    for el in list(header._element):
        header._element.remove(el)

    banner_path = Path(__file__).resolve().parents[1] / "web" / "assets" / "flying_header_banner.png"
    p_banner = header.add_paragraph()
    p_banner.paragraph_format.space_after = Pt(6)
    if banner_path.exists():
        run = p_banner.add_run()
        run.add_picture(str(banner_path), width=Cm(17.5))
    else:
        p_left = header.add_paragraph()
        _run(p_left, "G R U P O   F L Y I N G", size=9, color=COR_TEXTO_SOFT)
        p_logo = header.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_logo.paragraph_format.space_after = Pt(2)
        if LOGO_PATH.exists():
            run = p_logo.add_run()
            run.add_picture(str(LOGO_PATH), width=Cm(3.5))
        else:
            _run(p_logo, "FLYING studio", bold=True, size=14, color=COR_PRIMARIA)
        p_linha = header.add_paragraph()
        p_linha.paragraph_format.space_after = Pt(0)
        _add_horizontal_line(p_linha, color=HEX_PRIMARIA, size=6)


def _setup_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    for el in list(footer._element):
        footer._element.remove(el)

    p_linha = footer.add_paragraph()
    p_linha.paragraph_format.space_after = Pt(3)
    _add_horizontal_line(p_linha, color=HEX_PRIMARIA, size=6)
    # Hack: a linha é "above" → faço um texto vazio e linha em cima
    # Reverte: usa borda superior em vez de inferior
    p_pr = p_linha._p.get_or_add_pPr()
    p_bdr_existing = p_pr.find(qn("w:pBdr"))
    if p_bdr_existing is not None:
        p_pr.remove(p_bdr_existing)
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), HEX_PRIMARIA)
    p_bdr.append(top)
    p_pr.append(p_bdr)

    p_marcas = footer.add_paragraph()
    p_marcas.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_marcas.paragraph_format.space_after = Pt(4)
    _run(p_marcas, "FLYING studio", bold=True, size=8, color=COR_PRIMARIA)
    _run(p_marcas, "  |  www.flyingstudio.com.br     ", size=8, color=COR_TEXTO_SOFT)
    _run(p_marcas, "RINNO FILMS", bold=True, size=8, color=COR_TEXTO)
    _run(p_marcas, "  |  www.rinnofilms.com.br     ", size=8, color=COR_TEXTO_SOFT)
    _run(p_marcas, "nid.studio", bold=True, size=8, color=COR_TEXTO)
    _run(p_marcas, "  |  www.nidstudio.com.br", size=8, color=COR_TEXTO_SOFT)

    p_end = footer.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_after = Pt(0)
    _run(
        p_end,
        "AV. Engenheiro Luis Carlos Berrini, 936 - 7 andar - Itaim Bibi, São Paulo - SP, 04571-000  |  Telefone - (11) 2351-4138",
        size=7,
        color=COR_TEXTO_SOFT,
    )


# ---------- componentes ----------


def _caixa_capa(doc, *, cliente, qtd_img, valor_bruto, valor_final, desconto_pct, desconto_label, qtd_extras=0):
    """Tabela 2-colunas com fundo roxo dando o resumo da proposta."""
    linhas = [
        ("CLIENTE", cliente["empresa"].upper(), False),
        ("PROJETO", cliente["ref"].upper(), False),
        ("AOS CUIDADOS DE", cliente["contato"].upper(), False),
    ]
    if qtd_img:
        linhas.append(("IMAGENS", f"{qtd_img} unidades", False))
    if qtd_extras:
        linhas.append(("SERVIÇOS EXTRAS", f"{qtd_extras} {'item' if qtd_extras == 1 else 'itens'}", False))
    if desconto_pct > 0:
        linhas.append(("VALOR BRUTO", _brl(valor_bruto), False))
        linhas.append((f"DESCONTO ({desconto_label or f'{desconto_pct}%'})", "-" + _brl(valor_bruto - valor_final), False))
    linhas.append(("INVESTIMENTO", _brl(valor_final), True))

    tab = doc.add_table(rows=len(linhas), cols=2)
    tab.autofit = False
    tab.columns[0].width = Cm(6.0)
    tab.columns[1].width = Cm(10.0)

    for i, (rotulo, valor, destaque) in enumerate(linhas):
        c1 = tab.rows[i].cells[0]
        c2 = tab.rows[i].cells[1]
        c1.width = Cm(6.0)
        c2.width = Cm(10.0)

        _shade_cell(c1, HEX_PRIMARIA)
        _shade_cell(c2, HEX_PRIMARIA_DARK if destaque else HEX_PRIMARIA)
        _set_cell_no_borders(c1)
        _set_cell_no_borders(c2)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # padding via parágrafo
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(6)
        p1.paragraph_format.space_after = Pt(6)
        _run(p1, rotulo, bold=True, size=9, color=COR_BRANCO)

        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_before = Pt(6)
        p2.paragraph_format.space_after = Pt(6)
        _run(p2, valor, bold=True, size=16 if destaque else 11,
             color=COR_ACENTO if destaque else COR_BRANCO)


def _tabela_categoria(doc, numero, categoria, *, mostra_precos):
    """Tabela com header roxo, linhas zebra, subtotal escuro."""
    linhas_total = 1 + categoria.qtd + 1
    cols = 3 if mostra_precos else 2
    tab = doc.add_table(rows=linhas_total, cols=cols)
    tab.autofit = False
    if mostra_precos:
        tab.columns[0].width = Cm(2.0)
        tab.columns[1].width = Cm(10.5)
        tab.columns[2].width = Cm(3.5)
    else:
        tab.columns[0].width = Cm(2.5)
        tab.columns[1].width = Cm(13.5)

    # ===== Cabeçalho =====
    cab = tab.rows[0]
    cab_label_1 = "ITEM"
    cab_label_2 = "DESCRIÇÃO DO SERVIÇO"
    for c in cab.cells:
        _shade_cell(c, HEX_PRIMARIA)
        _set_cell_no_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cab.cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _run(p, cab_label_1, bold=True, size=9, color=COR_BRANCO)

    p = cab.cells[1].paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _run(p, cab_label_2, bold=True, size=9, color=COR_BRANCO)

    if mostra_precos:
        p = cab.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _run(p, "VALOR", bold=True, size=9, color=COR_BRANCO)

    # ===== Linhas =====
    for idx, item in enumerate(categoria.itens, start=1):
        row = tab.rows[idx]
        zebra = idx % 2 == 0  # idx começa em 1 → segunda linha visível recebe offwhite
        for c in row.cells:
            if zebra:
                _shade_cell(c, HEX_OFFWHITE)
            _set_cell_borders(c, bottom_color="E7E9EE", bottom_size=4)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        _run(p1, f"{numero}.{idx}", bold=True, size=10, color=COR_PRIMARIA_DARK)

        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after = Pt(3)
        _run(p2, item.descricao_normalizada, size=11, color=COR_TEXTO)

        if mostra_precos:
            p3 = row.cells[2].paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p3.paragraph_format.space_before = Pt(3)
            p3.paragraph_format.space_after = Pt(3)
            _run(p3, _brl(item.preco), bold=True, size=11, color=COR_TEXTO)

    # ===== Rodapé (subtotal) =====
    rod = tab.rows[-1]
    for c in rod.cells:
        _shade_cell(c, HEX_PRIMARIA_DARK)
        _set_cell_no_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = rod.cells[0].paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)
    _run(p1, str(categoria.qtd), bold=True, size=11, color=COR_BRANCO)

    if mostra_precos:
        p2 = rod.cells[1].paragraphs[0]
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        _run(p2, "Subtotal", bold=True, size=11, color=COR_BRANCO)
        p3 = rod.cells[2].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.paragraph_format.space_before = Pt(4)
        p3.paragraph_format.space_after = Pt(4)
        _run(p3, _brl(categoria.total), bold=True, size=11, color=COR_BRANCO)
    else:
        p2 = rod.cells[1].paragraphs[0]
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        _run(p2, f"Subtotal     {_brl(categoria.total)}", bold=True, size=11, color=COR_BRANCO)


# ---------- documento principal ----------


def _tabela_extra_subsecao(doc, numero: str, subsec: dict[str, Any]):
    """Tabela do estilo dos screenshots: cabeçalho roxo claro com título completo,
    sub-cabeçalho de colunas, itens numerados, valor total no rodapé."""
    itens = subsec.get("itens") or [subsec.get("rotulo_curto", "")]
    sem_preco = subsec.get("sem_preco", False)

    # 1 (cabeçalho de subseção) + 1 (cabec colunas) + N + 1 (rodapé)
    linhas_total = 3 + len(itens)
    tab = doc.add_table(rows=linhas_total, cols=2)
    tab.autofit = False
    tab.columns[0].width = Cm(2.5)
    tab.columns[1].width = Cm(13.5)

    # Linha 0: cabeçalho da subseção (lavanda claro com título completo)
    linha_cab = tab.rows[0]
    # Mescla as 2 células
    cell_a = linha_cab.cells[0]
    cell_b = linha_cab.cells[1]
    cell_merged = cell_a.merge(cell_b)
    _shade_cell(cell_merged, "EFEBFF")  # roxo bem claro
    # Bordas: linha roxa em cima e embaixo
    tc_pr = cell_merged._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for side, color in (("top", HEX_PRIMARIA), ("bottom", HEX_PRIMARIA)):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:color"), color)
        borders.append(b)
    for side in ("left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tc_pr.append(borders)
    cell_merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell_merged.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _run(p, f"{numero}  ", bold=True, size=11, color=COR_PRIMARIA_DARK)
    _run(p, subsec["rotulo_secao"], bold=True, size=11, color=COR_PRIMARIA_DARK)

    # Linha 1: cabeçalho de colunas (Itens / Descrição)
    linha_col = tab.rows[1]
    for c in linha_col.cells:
        _set_cell_borders(c, bottom_color="E7E9EE", bottom_size=4)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = linha_col.cells[0].paragraphs[0]
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after = Pt(3)
    _run(p1, "Itens", bold=True, size=9, color=COR_TEXTO_SOFT)
    p2 = linha_col.cells[1].paragraphs[0]
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after = Pt(3)
    _run(p2, "Descrição dos Serviços", bold=True, size=9, color=COR_TEXTO_SOFT)

    # Linhas de itens
    for idx, texto_item in enumerate(itens, start=1):
        row = tab.rows[1 + idx]
        for c in row.cells:
            _set_cell_borders(c, bottom_color="E7E9EE", bottom_size=4)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        _run(p1, f"{numero}.{idx}", bold=True, size=10, color=COR_PRIMARIA_DARK)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after = Pt(3)
        _run(p2, texto_item, size=11, color=COR_TEXTO)

    # Rodapé: contagem + Valor Total
    rod = tab.rows[-1]
    for c in rod.cells:
        _shade_cell(c, HEX_PRIMARIA_DARK)
        _set_cell_no_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = rod.cells[0].paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)
    _run(p1, str(len(itens)), bold=True, size=11, color=COR_BRANCO)
    p2 = rod.cells[1].paragraphs[0]
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    if sem_preco:
        _run(p2, "Valor Total — A DEFINIR", bold=True, size=11, color=COR_BRANCO)
    else:
        # Valor total na direita: usa um tab stop simulando alinhamento
        _run(p2, "Valor Total", bold=True, size=11, color=COR_BRANCO)
        # Tab + valor
        from docx.shared import Twips
        from docx.enum.text import WD_TAB_ALIGNMENT
        tab_stops = p2.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(13.0), WD_TAB_ALIGNMENT.RIGHT)
        run = p2.add_run("\t" + _brl(subsec["preco"]))
        run.font.name = FONTE
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = COR_BRANCO


def gerar_docx(
    *,
    cliente: dict[str, str],
    orc: Orcamento,
    saida: Path,
    data: _dt.date | None = None,
    mostra_precos_individuais: bool = False,
    forma_pagamento: list[dict[str, Any]] | None = None,
    prazos: dict[str, str] | None = None,
    extras: list[dict[str, Any]] | None = None,
    creditos: list[dict[str, Any]] | None = None,
    desconto_label: str | None = None,
    extras_estruturados: dict[str, Any] | None = None,
    usar_timbrado: bool = True,
    timbrado: Path | None = None,
) -> Path:
    data = data or _dt.date.today()
    tpl = resolver_timbrado(timbrado) if usar_timbrado else None
    if tpl:
        doc = Document(str(tpl))
        _limpar_corpo(doc)
    else:
        doc = Document()
        secao = doc.sections[0]
        secao.left_margin = Cm(2.5)
        secao.right_margin = Cm(2.5)
        secao.top_margin = Cm(3.0)
        secao.bottom_margin = Cm(2.3)
        secao.header_distance = Cm(1.2)
        secao.footer_distance = Cm(1.0)
        _setup_header(doc)
        _setup_footer(doc)

    subtotal_imagens = orc.subtotal
    total_extras = (extras_estruturados or {}).get("total", 0)
    qtd_extras = (extras_estruturados or {}).get("qtd", 0)
    qtd_imagens = orc.total_imagens
    subtotal = subtotal_imagens + total_extras
    desconto_valor = subtotal * (orc.desconto_pct / 100.0)
    valor_final = subtotal - desconto_valor

    # ===== CAPA =====
    _add_par(doc, "PROPOSTA COMERCIAL", bold=True, size=28, color=COR_PRIMARIA, space_before=18, space_after=2)
    _add_par(doc, "Imagens, Filmes e Tecnologias 3D", size=13, color=COR_TEXTO_SOFT, space_after=12)
    _add_par(doc, _data_extenso(data).upper(), size=9, color=COR_TEXTO_SOFT, space_after=24)

    _caixa_capa(
        doc,
        cliente=cliente,
        qtd_img=qtd_imagens,
        qtd_extras=qtd_extras,
        valor_bruto=subtotal,
        valor_final=valor_final,
        desconto_pct=orc.desconto_pct,
        desconto_label=desconto_label,
    )
    _add_par(doc, "", space_after=4)
    _add_par(doc, f"Por extenso: {_extenso(valor_final)}.", italic=True, size=9, color=COR_TEXTO_SOFT,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # quebra de página
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # ===== APRESENTAÇÃO =====
    _add_titulo_secao(doc, "01", "Apresentação")
    _add_par(doc,
             "A Flying Studio presta serviços de computação gráfica e tecnologias que se aplicam aos "
             "lançamentos imobiliários e remanescentes. Em nosso atendimento diário, desenvolvemos laços "
             "com projeto e auxiliamos em layout, estudos de projetos e fachadas, de decoração e "
             "paisagismo de acordo com cada necessidade.",
             color=COR_TEXTO_SOFT, space_after=6)
    _add_par(doc, "Para projetos de arquitetura, decoração e paisagismo, consulte a NID STUDIO.",
             italic=True, color=COR_TEXTO_SOFT, space_after=18)

    # ===== ITENS =====
    _add_titulo_secao(doc, "02", "Itens a Serem Desenvolvidos")

    secao_num = 0
    if orc.externas.qtd:
        secao_num += 1
        _add_par(doc, "ILUSTRAÇÕES EXTERNAS", bold=True, size=11, color=COR_PRIMARIA_DARK, space_before=8, space_after=4)
        _tabela_categoria(doc, f"2.{secao_num}", orc.externas, mostra_precos=mostra_precos_individuais)
        _add_par(doc, "", space_after=8)
    if orc.internas.qtd:
        secao_num += 1
        _add_par(doc, "ILUSTRAÇÕES INTERNAS", bold=True, size=11, color=COR_PRIMARIA_DARK, space_before=8, space_after=4)
        _tabela_categoria(doc, f"2.{secao_num}", orc.internas, mostra_precos=mostra_precos_individuais)
        _add_par(doc, "", space_after=8)
    if orc.plantas.qtd:
        secao_num += 1
        _add_par(doc, "PLANTAS HUMANIZADAS", bold=True, size=11, color=COR_PRIMARIA_DARK, space_before=8, space_after=4)
        _tabela_categoria(doc, f"2.{secao_num}", orc.plantas, mostra_precos=mostra_precos_individuais)
        _add_par(doc, "", space_after=8)

    # ===== EXTRAS (Tour Virtual / Filmes / Apps / Maquete / Drone / Estudo Fachada) =====
    if extras_estruturados and extras_estruturados.get("qtd", 0) > 0:
        for grupo_chave in ("tour_virtual", "filmes", "apps", "maquete", "drone", "estudo_fachada", "diversos"):
            grupo = extras_estruturados.get(grupo_chave) or {}
            for sub in grupo.get("subsecoes", []):
                secao_num += 1
                _tabela_extra_subsecao(doc, f"2.{secao_num}", sub)
                _add_par(doc, "", space_after=8)

    # Totais inline
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    if qtd_imagens:
        _run(p, "Imagens: ", color=COR_TEXTO_SOFT)
        _run(p, str(qtd_imagens), bold=True)
    if qtd_extras:
        if qtd_imagens:
            _run(p, "    ·    ", color=COR_TEXTO_SOFT)
        _run(p, "Serviços extras: ", color=COR_TEXTO_SOFT)
        _run(p, str(qtd_extras), bold=True)
    _run(p, "    ·    Valor bruto: ", color=COR_TEXTO_SOFT)
    _run(p, _brl(subtotal), bold=True)

    if orc.desconto_pct > 0:
        rotulo = desconto_label or f"{orc.desconto_pct}% de Desconto"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, "Desconto aplicado: ", color=COR_TEXTO_SOFT)
        _run(p, rotulo, bold=True)
        _run(p, "    ·    Valor do desconto: ", color=COR_TEXTO_SOFT)
        _run(p, "-" + _brl(desconto_valor), bold=True, color=COR_PRIMARIA_DARK)

    _add_par(doc, "", space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "INVESTIMENTO TOTAL  ", bold=True, size=14, color=COR_PRIMARIA)
    _run(p, _brl(valor_final), bold=True, size=20, color=COR_PRIMARIA_DARK)
    _add_par(doc, f"({_extenso(valor_final)})", italic=True, color=COR_TEXTO_SOFT, space_after=18)

    if extras:
        _add_par(doc, "EXTRAS / FILMES", bold=True, size=11, color=COR_PRIMARIA_DARK, space_before=8, space_after=4)
        for ex in extras:
            cortesia = " (CORTESIA)" if ex.get("cortesia") else ""
            preco = "" if ex.get("cortesia") else f" — {_brl(ex['preco'])}"
            _add_bullet(doc, f"{ex['descricao']}{preco}{cortesia}")
        _add_par(doc, "", space_after=8)

    # ===== nova página =====
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # ===== FORMA DE PAGAMENTO =====
    _add_titulo_secao(doc, "03", "Forma de Pagamento")
    fp = forma_pagamento or [
        {"percentual": 50, "marco": "Na aprovação desta Proposta"},
        {"percentual": 25, "marco": "Envio dos Shades"},
        {"percentual": 25, "marco": "Envio HR — Imagens finais"},
    ]
    for parc in fp:
        v = valor_final * (parc["percentual"] / 100.0)
        _add_bullet(doc, parc["marco"], label=f"{parc['percentual']}%  ({_brl(v)})")
    _add_par(doc, "", space_after=12)

    # ===== PRAZOS =====
    _add_titulo_secao(doc, "04", "Prazos de Entrega")
    pr = prazos or {
        "shades": "20 (Vinte) dias",
        "primeiro_tiro": "15 (Quinze) dias após a aprovação dos Shades",
        "revisoes": "10 (Dez) dias para contemplar e enviar novos tiros",
    }
    _add_bullet(doc, pr["shades"], label="Shades")
    _add_bullet(doc, pr["primeiro_tiro"], label="1º Tiro de Apresentação")
    _add_bullet(doc, pr["revisoes"], label="Revisões")
    _add_par(doc,
             "Os prazos passam a contar após o recebimento de todos os projetos, informações e aprovações "
             "de etapas para o desenvolvimento de cada item. Não iniciamos os trabalhos sem o DWG e as "
             "aprovações necessárias desta proposta.",
             italic=True, size=9, color=COR_TEXTO_SOFT, space_before=6, space_after=18)

    # ===== nova página =====
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # ===== MATERIAIS NECESSÁRIOS =====
    _add_titulo_secao(doc, "05", "Materiais Necessários")
    _add_bullet(doc, "Plantas · Elevação da Fachada · Estudo de Cores da Fachada · Cortes.", label="Arquitetura")
    _add_bullet(doc, "Implantação · Detalhamentos · Especificação de Revestimentos · Estudo de Vegetação com Especificação de Espécies · Referências do Mobiliário.", label="Paisagismo")
    _add_bullet(doc, "Plantas com Layout · Desenhos de Pisos · Elevações de Paredes · Especificações de Materiais · Projeto de Forro e Iluminação · Descrição ou book de mobiliários.", label="Decoração")
    _add_par(doc, "", space_after=12)

    # ===== CONSIDERAÇÕES =====
    _add_titulo_secao(doc, "06", "Considerações")
    consideracoes = [
        ("Etapas e Tiros de Aprovação", "Esta proposta contempla o envio inicial do tiro de Shade, seguido do tiro de apresentação denominado \u201CR00\u201D. Estão inclusas no escopo 03 (três) rodadas de revisões, denominadas \u201CR01\u201D, \u201CR02\u201D e \u201CR03\u201D, culminando na entrega final denominada \u201CHR\u201D (High Resolution)."),
        ("Ajustes Finos e Adicionais", "A partir do tiro \u201CR00\u201D, as rodadas seguintes consistem exclusivamente em ajustes finos. A partir de um eventual quarto tiro de apresentação (\u201CR04\u201D), será cobrado um adicional de 25% do valor da imagem por tiro extra solicitado, bem como quaisquer tiros adicionais solicitados após a entrega do HR."),
        ("Plataforma Oficial de Revisão", "Para garantir a organização, a agilidade e a precisão técnica das refações, todo o processo de feedback, comentários e aprovações (filmes e imagens 3D) será realizado exclusivamente através do software Frame.io/Adobe."),
        ("Mecânica de Apontamentos", "A Contratada fornecerá à Contratante um link de acesso seguro à plataforma. Pelo Frame.io/Adobe, o cliente poderá inserir comentários, desenhar marcações, anexar informações (pdf, foto, dwg, etc.) e solicitar ajustes exatamente no frame do vídeo ou no ponto específico da imagem estática que deseja alterar."),
        ("Alterações de Projeto", "Quaisquer alterações nos projetos originais (sejam de design de interiores, arquitetônico ou paisagismo) fornecidos inicialmente implicam em cobranças extras de modelagem, que serão orçadas e aprovadas em comum acordo."),
        ("Refação e Remodelagem", "Havendo mudanças significativas no projeto que resultem na perda de até 50% da imagem já construída, o trabalho será considerado e cobrado como uma imagem nova."),
        ("Paralisação do Projeto", "Em caso de paralisação total ou parcial do escopo por um período de até 60 (sessenta) dias, deverá ser feito o acerto financeiro imediato das etapas já executadas. Considera-se que cada tiro enviado após a aprovação do R00 corresponde a 25% do valor total da imagem."),
        ("Cancelamento", "Em caso de descontinuidade e cancelamento do produto ou lançamento por qualquer motivo por parte da Contratante, considera-se justa e devida a quitação integral do saldo previsto nesta proposta."),
        ("Direitos de Uso", "A Contratada cede à Contratante os direitos de uso das imagens produzidas para uso promocional em todo o seu material publicitário, única e exclusivamente vinculadas ao empreendimento contratado, não havendo débitos/atrasos financeiros."),
    ]
    for titulo, texto in consideracoes:
        _add_bullet(doc, texto, label=titulo)
    _add_par(doc, "", space_after=12)

    # ===== ENTREGA FINAL =====
    _add_titulo_secao(doc, "07", "Entrega Final")
    entregas = [
        ("Formato e Envio", "Todo o material finalizado será enviado digitalmente via servidor FTP, link seguro para download ou cadastrados no Frame.io/Adobe."),
        ("Resolução das Imagens Estáticas", "As imagens finais (\u201CHR\u201D) serão entregues com 6000px no lado maior a 300dpi. Após a entrega do HR, o projeto é considerado concluído. Caso surja a necessidade de novas configurações nessa etapa, ficamos à disposição para avaliar e orçar como um novo serviço."),
        ("Impressão de até 1m", "Caso a Contratante necessite de imagens configuradas para impressões de até 1 (um) metro, a solicitação deve ser feita com antecedência à renderização final, sem custo adicional."),
        ("Impressão acima de 1m", "Para outdoors ou grandes painéis (acima de 1m), favor consultar previamente os valores adicionais de render — custo estimado de 20% do valor da imagem."),
        ("Animações / Filmes", "Os passeios virtuais e filmes integrados serão entregues em Full HD a 30 FPS, ou propostas via RINNO FILMS, consultar."),
    ]
    for titulo, texto in entregas:
        _add_bullet(doc, texto, label=titulo)
    _add_par(doc, "", space_after=18)

    # ===== ASSINATURA =====
    _add_par(doc, f"São Paulo, {_data_extenso(data)}.", color=COR_TEXTO_SOFT, space_after=18)
    _add_par(doc, "De acordo,", space_after=24)
    _add_par(doc, "____________________________________________________", color=COR_TEXTO_SOFT, space_after=2)
    _add_par(doc, cliente["empresa"].upper(), bold=True, size=12, color=COR_PRIMARIA, space_after=2)
    _add_par(doc, f"A/C: {cliente['contato']}", color=COR_TEXTO_SOFT, size=9)

    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)
    return saida
